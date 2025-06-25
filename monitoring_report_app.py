import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from fpdf import FPDF
import matplotlib.pyplot as plt
from datetime import datetime
from PIL import Image
import base64
import os

st.set_page_config(page_title="E&S Monitoring Report Generator", layout="wide")

# Login Section
st.sidebar.header("🔐 Login")
def check_login():
    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")
    if username == "admin" and password == "yourpassword":
        return True
    else:
        st.sidebar.warning("Incorrect credentials")
        return False

if not check_login():
    st.stop()

st.title("Environmental & Social Monitoring Report Generator")

st.sidebar.header("Upload Section")

uploaded_cap = st.sidebar.file_uploader("Upload Corrective Action Plan (CAP) - Word or Excel", type=["docx", "xlsx"])
uploaded_esap = st.sidebar.file_uploader("Upload ESAP - Word or Excel", type=["docx", "xlsx"])
uploaded_photos = st.sidebar.file_uploader("Upload Site Visit Photos (JPG/PNG)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

# Helper to read Word table
def read_docx_table(uploaded_file):
    doc = Document(uploaded_file)
    table = doc.tables[0]
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return pd.DataFrame(data[1:], columns=data[0])

# Helper to write PDF
def generate_pdf(df, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, title, ln=True, align='C')
    pdf.set_font("Arial", '', 10)
    col_widths = [180 / len(df.columns)] * len(df.columns)
    pdf.set_fill_color(200, 220, 255)
    for col in df.columns:
        pdf.cell(col_widths[0], 10, col, border=1, fill=True)
    pdf.ln()
    for _, row in df.iterrows():
        for val in row:
            pdf.cell(col_widths[0], 10, str(val), border=1)
        pdf.ln()
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

# Summary generation
def generate_summary(df, title):
    total = len(df)
    completed = len(df[df['Status'].str.lower() == 'completed']) if 'Status' in df.columns else 0
    pending = total - completed
    progress = f"{completed}/{total} completed ({(completed/total*100 if total else 0):.1f}%)"
    st.container().markdown(f"**{title} Summary:**\n\n- Total Items: {total}\n- Completed: {completed}\n- Pending: {pending}\n- Progress: {progress}")
    fig, ax = plt.subplots()
    ax.pie([completed, pending], labels=['Completed', 'Pending'], autopct='%1.1f%%', colors=['#4CAF50','#FFC107'])
    ax.set_title(f"{title} Completion Status")
    st.pyplot(fig)

# Timeline Chart
def generate_timeline(df, title):
    if 'Due Date' in df.columns and 'Topic' in df.columns:
        try:
            df['Due Date'] = pd.to_datetime(df['Due Date'], errors='coerce')
            df = df.dropna(subset=['Due Date'])
            df_sorted = df.sort_values('Due Date')
            fig, ax = plt.subplots(figsize=(10, 4))
            ax.barh(df_sorted['Topic'], df_sorted['Due Date'].map(datetime.toordinal), color='#2196F3')
            ax.set_yticks(range(len(df_sorted)))
            ax.set_yticklabels(df_sorted['Topic'])
            ax.set_title(f"{title} Timeline by Due Date")
            st.pyplot(fig)
        except Exception as e:
            st.warning(f"Could not render timeline chart: {e}")

# Executive Summary Generator
def generate_exec_summary(cap_df, esap_df):
    cap_total = len(cap_df)
    cap_completed = len(cap_df[cap_df['Status'].str.lower() == 'completed']) if 'Status' in cap_df.columns else 0
    esap_total = len(esap_df)
    esap_completed = len(esap_df[esap_df['Status'].str.lower() == 'completed']) if 'Status' in esap_df.columns else 0
    return f"During this reporting period, progress was observed in both the Corrective Action Plan (CAP) and the Environmental and Social Action Plan (ESAP). For the CAP, {cap_completed} of {cap_total} actions ({(cap_completed/cap_total*100 if cap_total else 0):.1f}%) have been completed. The ESAP recorded {esap_completed} out of {esap_total} ({(esap_completed/esap_total*100 if esap_total else 0):.1f}%) actions completed. Continuous monitoring and strategic coordination remain essential to address the remaining gaps."

# Recommendations Generator
def generate_recommendations(df):
    if 'Status' in df.columns and 'Risk Level' in df.columns:
        high_risk_pending = df[(df['Status'].str.lower() != 'completed') & (df['Risk Level'].str.lower() == 'high')]
        med_risk_pending = df[(df['Status'].str.lower() != 'completed') & (df['Risk Level'].str.lower() == 'medium')]
        recs = ["**Conclusions and Recommendations:**"]
        if not high_risk_pending.empty:
            recs.append(f"- There are {len(high_risk_pending)} high-risk actions still pending. These must be prioritized and closed promptly to avoid potential compliance gaps or operational risks.")
        if not med_risk_pending.empty:
            recs.append(f"- There are {len(med_risk_pending)} medium-risk actions pending. These should be addressed in the upcoming period through adequate resource allocation and oversight.")
        if not recs[1:]:
            recs.append("- No high or medium risk items remain open. Maintain current monitoring and reporting cadence.")
        return '\n'.join(recs)
    return "Risk analysis and recommendations will appear here once 'Status' and 'Risk Level' columns are available."

# Table of Contents
st.markdown("## \U0001F4C1 Table of Contents")
toc_items = [
    "1. Executive Summary",
    "2. Introduction",
    "3. Project Description and Status",
    "3.1 Project Description",
    "3.3 Project Participants",
    "4. Project Permits and Authorizations",
    "5. Environmental, Health, Safety and Social Performance Review",
    "6. Corrective Action Plans",
    "7. Environmental and Social Action Plan – SIBA Thermoelectric Project",
    "8. Site Visit and Stakeholder Engagement",
    "9. Conclusions and Recommendations",
    "10. Annex 1"
]
for item in toc_items:
    st.markdown(f"- {item}")

# Executive Summary Section
st.subheader("1. Executive Summary")
cap_df, esap_df = pd.DataFrame(), pd.DataFrame()

# CAP Section
if uploaded_cap:
    st.subheader("6. Corrective Action Plan (CAP) Table")
    if uploaded_cap.name.endswith(".docx"):
        cap_df = read_docx_table(uploaded_cap)
    else:
        cap_df = pd.read_excel(uploaded_cap)

    if 'Risk Level' in cap_df.columns:
        risk_options = cap_df['Risk Level'].unique().tolist()
        selected_risks = st.multiselect("Filter CAP by Risk Level", risk_options, default=risk_options)
        cap_df = cap_df[cap_df['Risk Level'].isin(selected_risks)]

    generate_summary(cap_df, "CAP")
    generate_timeline(cap_df, "CAP")
    cap_df_edit = st.data_editor(cap_df, use_container_width=True, num_rows="dynamic", key="cap_editor")
else:
    st.info("Please upload a CAP file.")

# ESAP Section
if uploaded_esap:
    st.subheader("7. Environmental and Social Action Plan (ESAP) Table")
    if uploaded_esap.name.endswith(".docx"):
        esap_df = read_docx_table(uploaded_esap)
    else:
        esap_df = pd.read_excel(uploaded_esap)

    if 'Risk Level' in esap_df.columns:
        risk_options = esap_df['Risk Level'].unique().tolist()
        selected_risks = st.multiselect("Filter ESAP by Risk Level", risk_options, default=risk_options, key='esap')
        esap_df = esap_df[esap_df['Risk Level'].isin(selected_risks)]

    generate_summary(esap_df, "ESAP")
    generate_timeline(esap_df, "ESAP")
    esap_df_edit = st.data_editor(esap_df, use_container_width=True, num_rows="dynamic", key="esap_editor")
else:
    st.info("Please upload an ESAP file.")

# Site Visit Section
st.subheader("8. Site Visit and Stakeholder Engagement")
st.text_area("Site Visit Observations and Findings", "")
st.text_area("Stakeholder Engagement Activities", "")
if uploaded_photos:
    for img_file in uploaded_photos:
        st.image(Image.open(img_file), caption=img_file.name, use_column_width=True)

# Historical Comparison Section
st.subheader("\U0001F4CA Historical Comparison")
st.info("Future enhancement: Load and compare previous reports here to track progress over time.")

# Auto Executive Summary Output
if not cap_df.empty and not esap_df.empty:
    st.markdown("---")
    st.markdown("### Auto-Generated Executive Summary")
    st.write(generate_exec_summary(cap_df, esap_df))
    st.markdown("### Auto-Generated Conclusions and Recommendations")
    st.write(generate_recommendations(pd.concat([cap_df, esap_df], ignore_index=True)))

# Allow Download of Tables as Word

def to_word(df, title):
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.sidebar.header("Download Reports")
if uploaded_cap:
    cap_docx = to_word(cap_df_edit, "Corrective Action Plan (CAP)")
    st.sidebar.download_button("Download CAP (Word)", cap_docx, file_name="CAP_Report.docx")
    cap_pdf = generate_pdf(cap_df_edit, "Corrective Action Plan (CAP)")
    st.sidebar.download_button("Download CAP (PDF)", cap_pdf, file_name="CAP_Report.pdf")

if uploaded_esap:
    esap_docx = to_word(esap_df_edit, "Environmental and Social Action Plan (ESAP)")
    st.sidebar.download_button("Download ESAP (Word)", esap_docx, file_name="ESAP_Report.docx")
    esap_pdf = generate_pdf(esap_df_edit, "Environmental and Social Action Plan (ESAP)")
    st.sidebar.download_button("Download ESAP (PDF)", esap_pdf, file_name="ESAP_Report.pdf")


